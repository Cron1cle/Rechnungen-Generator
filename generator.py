import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os
import glob
import subprocess

doc = docx.Document('template.docx')
# Utils
class color:
   PURPLE = '\033[95m'
   CYAN = '\033[96m'
   DARKCYAN = '\033[36m'
   BLUE = '\033[94m'
   GREEN = '\033[92m'
   YELLOW = '\033[93m'
   RED = '\033[91m'
   BOLD = '\033[1m'
   UNDERLINE = '\033[4m'
   END = '\033[0m'

title_text = color.UNDERLINE + color.RED + 'Invoice ' + color. YELLOW + 'Generator '+ color.GREEN+'v2'

for _ in range(100):
    print()

print(color.BLUE + f"""
                ▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓
                ▓▓▓                                           ▓▓▓
               ▓▓▓  ▄██                   ▄██ █                ▓▓▓
               ▓▓▓ ██▀  █  ▄█▄ █▄▄▄ █ ▄██ █   █   ▄█▄ ▄█▄  ▄▀▀ ▓▓▓
               ▓▓▓ ██   ██ █ █ █▀▀█   █   ▀█▄ ███ █ █ █  █ █   ▓▓▓
               ▓▓▓ ██▄  █  █ █ █  █ █ █     █ █ █ █ █ █▀▀  ▀▀█ ▓▓▓
               ▓▓▓  ▀██ █  ▀█▀ █  █ █ ▀██ ██▀ █ █ ▀█▀ ▀▄▄  ▄▄▀ ▓▓▓
                ▓▓▓                                           ▓▓▓
                ▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓
      
                            {title_text}
""" + color.END)




path = os.path.dirname(os.path.abspath(__file__))
def get_last_invoice_name():
    try:
        pdf_files = glob.glob(os.path.join(path, '*.pdf'))
        newest_pdf = max(pdf_files, key=os.path.getctime)
        return os.path.basename(newest_pdf)
    except Exception as e:
        return str(e)


# Rechnungsadresse
print(color.UNDERLINE + color.BOLD + color.PURPLE + 'Rechnungsadresse:' + color.END)
name = input('Name: ')
address = input('Adresse: ')
zipcode = int(input('PLZ: '))
city = input('Stadt: ')
country = input('Land: ')
vatid = input('VAT-Nummer: ')
print()

# Lieferadresse
d_name = name
d_address = address
d_zipcode = zipcode
d_city = city
d_country = country
d_vatid = vatid

address_equal = input('Weicht die Lieferadresse von der Rechnungsadresse ab? (J/N): ')
if address_equal.lower() == 'j' or address_equal.lower() == 'ja' or address_equal.lower() == 'yes' or address_equal.lower() == 'y':
    print(color.UNDERLINE + color.PURPLE + color.BOLD + 'Rechnungsadresse:' + color.END)
    d_name = input('Name: ')
    d_address = input('Adresse: ')
    d_zipcode = int(input('PLZ: '))
    d_city = input('Stadt: ')
    d_country = input('Land: ')
    d_vatid = input('VAT-Nummer: ')
else:
    print(color.UNDERLINE + color.PURPLE + color.BOLD + 'Lieferadresse:' + color.END)
    print('Name: ' + d_name)
    print('Adresse: ' + d_address)
    print('PLZ: ' + str(d_zipcode))
    print('Stadt: ' + d_city)
    print('Land: ' + d_country)
    print('VAT-Nummer: ' + d_vatid)
print()


# Rechnungsdaten
print(color.UNDERLINE + color.PURPLE + color.BOLD + 'Rechnungsdaten:' + color.END)
date_default = datetime.datetime.now().strftime("%d.%m.%Y")
invoice_n_default = str(datetime.datetime.now().year)+'0001'
if get_last_invoice_name().startswith('Rechnung-'):
    invoice_n_default = str(int(get_last_invoice_name().split('-')[1].split('.')[0]) + 1)
customer_n_default = invoice_n_default[-4:]

date = input(f'Datum ({date_default}): ')
if date == '':
    date = date_default
invoice_n = input(f'Rechnungsnummer ({invoice_n_default}): ')
if invoice_n == '':
    invoice_n = invoice_n_default
customer_n = input(f'Kundennummer ({customer_n_default}): ')
if customer_n == '':
    customer_n = customer_n_default


# Bankverbindung
bankname = 'BANKNAME'
blz = 'BANKLEITZAHL'
kto = 'KONTONUMMER'
iban = 'IBAN'

paypal_account = 'PAYPAL'





change_bank = input('Bankverbindung ändern? (J/N):')
print()
print(color.UNDERLINE + color.PURPLE + color.BOLD + 'Bankverbindung:' + color.END)
if change_bank.lower() == 'ja' or change_bank.lower() == 'j' or change_bank.lower() == 'yes' or change_bank.lower() == 'y':
    bankname = input('Bankname: ')
    blz = input('BLZ: ')
    kto = input('Kontonummer: ')
    iban = input('IBAN: ')
else:
    choose_bank = input('Welche Bankverbingdung möchtest du? (C/S):'+ color.END)
    print()
    if choose_bank.lower() == 's':
        print('Bankname: ' + bankname)
        print('BLZ: ' + blz)
        print('Kontonummer: ' + kto)
        print('IBAN: ' + iban)
    else: 
        print()      





table = doc.tables[0]

product_amount = int(input('Gib die Anzahl der unterschiedlich gekauften Produkte an: '))
for _ in range(product_amount):
    new_row = table.add_row()
    
new_row._tr.addnext(table.rows[1]._tr)


print()
print('—————————————————————————————————————')
p_sum = 0
for i in range(product_amount):
    print(color.UNDERLINE + color.BOLD + color.CYAN + 'Produkt ' + str(i+1) + color.END)
    p_amount = input('Anzahl (eg. 2): ')
    p_num = input('Artikelnummer (eg. 553558-093): ')
    p_name = input('Bezeichnung (eg. Air Jordan 1 Low): ')
    p_singleprice = float(input('Einzelpreis (eg. 129.99): '))
    p_price = p_singleprice * float(p_amount)

    p_sum += p_price

    table.rows[i+1].cells[0].text = str(p_amount)  + 'x'
    table.rows[i+1].cells[1].text = 'Stk.'
    table.rows[i+1].cells[2].text = p_num
    table.rows[i+1].cells[3].text = p_name

    table.rows[i+1].cells[4].text = str('%05.2f' % p_singleprice).replace('.', ',') + '€'
    table.rows[i+1].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table.rows[i+1].cells[5].text = str('%05.2f' % p_price).replace('.', ',')  + '€'
    table.rows[i+1].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if i < product_amount-1:
        print('—————————————————————————————————————')

table.rows[-1].cells[-1].text = str('%05.2f' % p_sum).replace('.', ',')  + '€'
sum_paragraph = table.rows[-1].cells[-1].paragraphs[0]
sum_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in sum_paragraph.runs:
    run.bold = True

# Iteriere durch alle Absätze im Dokument
for para in doc.paragraphs:
    if "{" in para.text:
        for run in para.runs:
            if "{" in run.text:
                run.text = run.text.replace("{NAME}", name)
                run.text = run.text.replace("{ADDY}", address)
                run.text = run.text.replace("{ZIPCODE}", str(zipcode))
                run.text = run.text.replace("{CITY}", city)
                run.text = run.text.replace("{COUNTRY}", country)
                run.text = run.text.replace("{VATID}", vatid)
                run.text = run.text.replace("{NAME_D}", d_name)
                run.text = run.text.replace("{ADDY_D}", d_address)
                run.text = run.text.replace("{ZIPCODE_D}", str(d_zipcode))
                run.text = run.text.replace("{CITY_D}", d_city)
                run.text = run.text.replace("{COUNTRY_D}", d_country)
                run.text = run.text.replace("{DATE}", date)
                run.text = run.text.replace("{INVOICE_N}", invoice_n)
                run.text = run.text.replace("{CUSTOMER_N}", customer_n)

# Footer
section = doc.sections[0]
footer = section.footer
footer.tables[0].rows[1].cells[0].paragraphs[0].text = bankname + ' (BLZ: ' + blz + ')'
footer.tables[0].rows[1].cells[0].paragraphs[1].text = 'Konto-Nummer: ' + kto
footer.tables[0].rows[1].cells[0].paragraphs[2].text = 'IBAN: ' + iban




print(color.YELLOW + color.BOLD + 'Rechnung wird generiert...')
doc.save(f'Rechnung-{invoice_n}.docx')
print('Konvertiere zu PDF...'+ color.END)
from docx2pdf import convert
convert(f'Rechnung-{invoice_n}.docx')
print()
print(color.GREEN + color.BOLD + 'Rechnung wurde erfolgreich generiert und konvertiert.' + color.END)

# Lösche das Word-Dokument
os.remove(f'Rechnung-{invoice_n}.docx')

# PDF öffnen
open_invoice = input('Rechnung öffnen? (J/N): ')
if open_invoice.lower() == 'ja' or open_invoice.lower() == 'j' or open_invoice.lower() == 'yes' or open_invoice.lower() == 'y':
    print(color.YELLOW + color.BOLD + 'Rechnung wird geöffnet...')
    subprocess.Popen(['Rechnung-' + invoice_n+'.pdf'],shell=True)
else:
    print(color.GREEN + color.BOLD + 'Rechnung gespeichert unter ' + str(path) + '\\' + invoice_n+'.pdf' +  + color.END)

